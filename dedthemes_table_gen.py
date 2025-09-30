from docx import Document
from docx.shared import Pt

# Create new Word document
doc = Document()

# Add table title in APA style (centered, italicized)
title = doc.add_paragraph("Table 1\n")
title.alignment = 1
run = title.add_run("Deductive Themes, Subthemes, and Sample Participant Quotes")
run.italic = True

# Prepare structured table (deductive theme, subtheme, quotes)
data = [
    ["Deductive Theme", "Subtheme", "Quotes"],
    ["General appraisal of the Box", "Box was aesthetically pleasing",
     "“The quality of these materials is really impressive. Just putting it in my hands and touching it. "
     "I'm like, “wow, it’s very high quality. Good job with that.” (P6)\n\n"
     "“Even the box [is really nice].” (P5)\n\n"
     "“We watched that video together, that was really nice. Overall, I would just say that it catered to a lot "
     "of different attention styles. If I was more auditory there was things I could listen to, if I liked to read "
     "it more there were things I could read, if I’m a more physical learner there are things I can manipulate, "
     "it was really interactive in a lot of different ways.” (P13)"],
    ["", "Box was relevant to work",
     "“This is a box for me. It's not a box for me to help someone else. In all seriousness, that's all I do the "
     "whole day. Every day I have a master's degree in helping other people. Everything you give me I'm going to try "
     "to implement in my practice. But the fact that this is for me, it's like a dream come true. Someone knows that "
     "there is secondary trauma in school counseling…!” (P9)\n\n"
     "“I think a lot of the stuff that I found in the School Counselor Box is very relevant to what I do on a daily "
     "basis with my students, like Kris was saying before…the post-it notes, everything like the research tools that I "
     "found in there to use just to reinforce with my students as they come into my office on a daily basis” (P2)\n\n"
     "“But yeah, I realized when I heard and saw the video: I was like, ‘oh, that's for us!’ I thought: ‘okay, this is neat!’ "
     "Cause normally, we don't have anything that's counselor focused, because we are the people that help students deal "
     "with trauma. And yet we, too, are affected by the same information that we hear. Even if indirectly, because it can "
     "be overwhelming and taxing to hear some of the things that you hear on a daily basis.” (P12)"],
    ["", "Box included more foundational content but was helpful",
     "“Well, I was aware of compassion fatigue and secondary trauma, but it was very good to actually have it broken down "
     "[through] the video, the different lessons, and resources. Because I think a lot of times we feel guilty trying to "
     "take care of ourselves. So yeah, I just needed permission [to engage in self-care], for someone to say “you all are "
     "going through a lot”.” (P4)\n\n"
     "“A lot of the stuff that they talk about [in the Box] I knew, but for me it was a reminder. Because, even though you "
     "know things, just having that reminder, having somebody to say it out loud [is helpful].” (P10)"],
    ["", "Box was personally validating",
     "“So it was very…. how can I say? It gave me hope. The box gave something for me. And I like these stickies, because… "
     "that's something that I can share with others… and it brings me joy.” (P9)"],
    ["Most or Least Helpful Items in the Box", "Most helpful",
     "(Book) “I started reading it, and I was like, wow! This is just for me as a support to help me as a counselor and other "
     "helpers within the field like teachers… it was really insightful, some of the things that I was reading, things that I "
     "found out about.” (P11)\n\n"
     "(Book) “The book is phenomenal. I think that is something that needs to be shared with everyone… It's just very applicable "
     "to what we're dealing with among our students, how we're trying to support them, and how that affects us, even when we realize "
     "we're just kind of numb to it.” (P6)\n\n"
     "(Sticky notes) “I actually gave one to our [behavioral health provider] today because she was having a rough morning.” (P4)\n\n"
     "(Sticky notes) “The post-it notes were very helpful. They’re easy, they’re to the point, and they’re great reminders.” (P1)\n\n"
     "(Video) “I thought the video explained it [vicarious trauma] very well. I remember the feeling of watching the video saying, "
     "“Okay, I know why this is created”. I liked watching the video, and I’m a person who usually skips videos! So I appreciate it.” (P9)\n\n"
     "(Video) “And as I was listening to the video, all I could think about is that I wanted to share this with some of my colleagues who "
     "are struggling particularly right now with things that's happening within the school system.” (P10)"],
    ["", "Least helpful",
     "“I wouldn't say that it wasn't relevant, but the [perspective] cards, initially those pictures didn’t really resonate with me.” (P11)\n\n"
     "“I probably wouldn't use that [peer-to-peer cards because] teachers do not really have a lot of time. I think it’s more important to "
     "just have a conversation.” (P4)"],
    ["Impact of Using the Box", "Increased self-reflection or self-awareness",
     "“I think for me it’s very hard for me to admit sometimes that I’m suffering or that I’m having a reaction because I’m so busy taking care "
     "of other people. This box for me helped me to at least acknowledge it and say it out loud..” (P10)\n\n"
     "“I would say that using the word trauma really helped. I wasn’t really prepared to apply what I learned with crisis into this role but I "
     "think it made me realize I’m going to have to do that. I’m really lucky to have the box early into my career because I can kind of set "
     "things right from the beginning.” (P13)\n\n"
     "“​​All these people are depending on me. What am I going to do? And I neglect me. If don’t get me together first, I can’t do anything else. "
     "So I think this is good to really remind yourself to check on you first.” (P12)"],
    ["", "Increased self-care",
     "“So this box actually helped because what I do is I take those sticky notes and I stick it on my computer. [When] all of this [stress] is "
     "going on, it feels like the space is going to burn down, but I need to focus on this for me. I need to do some deep breathing. I need to calm "
     "my mind before I extend myself and to be more useful to those around me.” (P11)\n\n"
     "“[I now] remind myself to really continue to do self-care, which for me means getting out and exercising and riding my bike. I had stopped doing "
     "that for a while, and I can’t tell you why, but I stopped.” (P10)\n\n"
     "“I have made more of an effort in terms of engaging in self care… and that was thanks to the box!” (P9)"],
    ["", "Reached out to colleagues to ask or offer support",
     "“I used it to bring up the conversation again to the [behavioral health provider]. And I will make sure that I incorporate it with our mental health "
     "team, the psychologists, and everyone else, because it's important and it's important to recognize it.” (P4)\n\n"
     "“Since I’ve had the box, it helped me to show even more how I care about other teachers… I have been [reaching out to show] care about how their "
     "families are doing, how our day [is going]… so the box has helped me work or get closer to other teachers, and to reinforce [to] them that what they "
     "do transmits to children and they need to be happy. We need to also look after ourselves.” (P5)"],
    ["", "Psychological benefits of understanding the experience of ‘trauma’",
     "“So I know what it is because I have seen it, and I have experienced it. But the terminology literally put a stamp on it for me. To put it under a label. "
     "So that was one major learning curve for me from this box.”..“So when this term was introduced.... And then I realize, this is something that I cannot ignore. "
     "I have accepted it, and I know that I have to pay more attention to self-care.” (P11)\n\n"
     "“Well now I have a word. To describe what I was feeling. I didn’t put the term with what I was feeling. (P12)”\n\n"
     "“I would say that using the word trauma really helped. I wasn’t really prepared to apply what I learned with crisis into this role but I think it made me realize "
     "I’m going to have to do that. I’m really lucky to have the box early into my career because I can kind of set things right from the beginning…… Yeah! That was one "
     "of those really helpful things. When I went through t with my supervisor, she was like, I've never really thought about that word trauma with what we go through sometimes.” (P13)"],
    ["Where School Counselors Get Information", "Educational program/master’s degree",
     "“The trainings that I had [on vicarious trauma], they were in my master's degree. I didn't get my counseling master's degree until 2014, so a lot of it is still fresh in my mind.” (P1)"],
    ["", "Professional organization",
     "Social Workers of America Conference (P9)\n\nNational Alliance on Mental Illness (P1)\n\nPREPaRE - Curriculum developed by the National Association of School Psychologists"],
    ["", "Personal research",
     "“I prefer Facebook, it's more pleasant, it has pictures… I see interesting topics from reliable sources I follow… But if I really want to study something and understand it, I have to go back to whatever materials I had before [from my masters degree] meaning peer reviewed ones.” (P10)"],
    ["Recommendations for Enhancing the Box", "",
     "Multiple participants would have liked some more suggestions on how to use the different materials with colleagues or students in different ways (beyond instructions for using the item for themselves)\n\n"
     "Participants with more exposure to vicarious trauma trainings or research would have liked to see more advanced content, but they weren’t exactly sure what to recommend\n\n"
     "A couple recommended adding content related to suicide or other mental health concerns that are associated with vicarious trauma among school professionals\n\n"
     "Some recommended more tailored activities depending on a person’s interests or for different points in the day (e.g., when at work or when they were driving home)\n\n"
     "Would like a more expansive resource library on the USB drive"]
]

# Build table with borders
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Add headers
hdr_cells = table.rows[0].cells
for i, val in enumerate(data[0]):
    hdr_cells[i].text = val

# Add all rows
for row in data[1:]:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val

# Add note at the bottom
note = doc.add_paragraph("\nNote. Participants are labeled with a P and a random identifying number to differentiate between individuals.")

# Save
file_path = "/mnt/data/Deductive_Themes_APA.docx"
doc.save(file_path)

file_path
