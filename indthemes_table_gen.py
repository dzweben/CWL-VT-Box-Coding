from docx import Document
from docx.shared import Pt

# Create new Word document
doc = Document()

# Set font to Times New Roman
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Add title
doc.add_heading("Table X\nInductive Themes with Sample Quotes", level=1)

# Create table with 3 columns
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Add headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Theme"
hdr_cells[1].text = "Subtheme"
hdr_cells[2].text = "Sample Quotes"

# Data (themes, subthemes, quotes with line breaks)
themes = [
    ("Intentions to use the Box with colleagues and students", "Thought Box was for students",
     "“In all honesty, I originally thought it was a box for tools for students… Cause normally, we don't have anything that's counselor focused… and yet we too are affected by the same information that we hear.” (P12)\n\n"
     "“Initially when I received that I thought that it would have been something that I could’ve used directly with my students.” (P11)\n\n"
     "“I think a lot of the stuff that I found in the School Counselor Box is very relevant to what I do on a daily basis with my students…everything like the research tools that I found in there to use just to reinforce with my students as they come into my office on a daily basis, have all been beneficial to the work that I do with my students."),
    
    ("Intentions to use the Box with colleagues and students", "Box may be helpful for teachers or other school professionals",
     "“I believe that all the information [in the Box] we kind of know, but it's always good to have a refresher and have an approach of how to do everything in little steps. Then we can help not only students but also our colleagues and teachers.” (P5)\n\n"
     "“It would be nice to share it [the Box] with teachers because they go through a lot of trauma themselves, and they don't know how to deal with it. And then we have to be there to listen to them.” (P1)\n\n"
     "“And as I was listening to the video, all I could think about is that I wanted to share this with some of my colleagues who are struggling particularly right now with things that's happening within the school system.” (P10)"),
    
    ("Challenges navigating mental health demands in the school system", "Empathy for vicarious trauma among teachers or other school professionals",
     "“You see it [vicarious trauma] in the teachers, too. It's very relevant. Sometimes it’s like, “why can't this teacher like get it together” or they're late or they're forgetful, but they're dealing with things too.” (P4)\n\n"
     "“I just got an email yesterday from a teacher that seemed really broken, so I'm trying my best to really support her even though I don't know her that well… I don't like seeing my teachers suffering...” (P3)\n\n"
     "“I’ve seen situations where educators are so burnt out that it starts to be the simple things they are forgetting. Like misplacing something important. So I definitely understand that part.” (P12)"),
    
    ("Challenges navigating mental health demands in the school system", "Burnout and challenges in meeting children’s and families’ needs",
     "“I have a caseload of approximately 250 to 300 students. It's a very large school. [School redacted] is approximately 1,400-1,500 students. It's very hard to provide that individualized attention to everyone who needs it.” (P2)\n\n"
     "“I would say that no one – at least that I’ve seen so far at least at my school – really acknowledges that most of the counselors are glued to their phones on the weekends. And they have to be.” (P13)\n\n"
     "“When I go home I'm exhausted. I am mentally drained, physically, everything. Because you're putting on a show, you're there for the kids, you're talking about the [mental health] lessons, and they're really into it, but it's so much. It's just heavy.” (P1)\n\n"
     "“I’m like yeah how can I get more hours in my day to see this student or that student, to get this project done. Sometimes I go without lunch. And I still don’t have enough time to get things done…And I’m like, okay, it's frustrating to me. Because I want to be there. I want to assist. But then I feel overwhelmed. And I feel like everything is just coming down on me. And we are the only school counselors, so it’s not as if there is someone else within my [department]. I mean we have our mental health worker and BHP, but there is no one else who is doing the same type of job that we are that we can go to and just let it all out” (P11)"),
    
    ("Challenges navigating mental health demands in the school system", "Impact of recent legislation",
     "“Policies change! And they don't always tell you until the last minute. So there's always this fear, at least for me, that I’m going to miss something… and I'm going to be held accountable for it.” (P9)\n\n"
     "“For me it's the state of Florida legislation. It's just every day I find something new.” [continued to discuss the inability to counsel students on certain high risk topics due to policy restrictions] (P10)\n\n"
     "“I feel like little bits and pieces [of new legislation] are being thrown at us. Not one whole set of instructions. It’s confusing. My mind is confused about it… it’s too many things, and they are not really giving us A to Z training about something. Just fill out this form. Fill out this form. Fill out this form.” (P7)"),
    
    ("Challenges navigating mental health demands in the school system", "Lingering impact of COVID on youth mental health",
     "“At first I thought ‘no, no it’s not COVID. COVID is not something going on right now. I mean it’s going on, but not like the past three years. This is the first year that I’m not thinking about COVID. But then I stopped and thought, oh wait, it is about COVID because, in my opinion, it caused a higher need for mental health services. I never had so many parents requesting support and school counseling for their kids like I had this year.” (P10)\n\n"
     "“[COVID] was a big toll on everyone emotionally. I think if I had something like this [Box] at that time it would’ve helped me a lot. Especially when we were all remote it was very challenging and stressful.” (P12)\n\n"
     "“We're still dealing with things post COVID. All the deaths that our kids and faculty and experienced. And we even experienced losing a faculty member – two of them actually. So it's just a lot.” (P4)"),
    
    ("Difficulty prioritizing self-care in the counseling role", "",
     "“It’s so true. We forget to take care of ourselves. The daily mantra of life, ‘I just have to do what I have to do’, or ‘someone needs me right now’ and you just go. But you don’t really take the time to eat, to think about what you need to get through your day. Maybe you had a rough night and you are coming into work and you have to put on a mask. Really, deep down inside, you need a moment, you need a space to debrief or to cry away from the view of others. To kind of gather yourself.” (P12)\n\n"
     "“It does help me, the box. But I think, again as I said, just going back to re-listen, to remind myself to really continue to do self-care. Which for me means getting out and exercising and riding my bike, because I had stopped doing that for a while. And I can’t tell you why. But I stopped. And I realized I wasn’t feeling as rested and relaxed when I stopped. So I think for me, going back and listening to the video will be very helpful…because I love to listen, and it just motivates me.” (P10)\n\n"
     "“[Counselors] don’t really have clear boundaries. Nobody acknowledges that. Or god forbid says not to do that.” (P13)"),
]

# Populate the table
for theme, subtheme, quotes in themes:
    row_cells = table.add_row().cells
    row_cells[0].text = theme
    row_cells[1].text = subtheme
    row_cells[2].text = quotes

# Save file
output_path = "/mnt/data/VT_BOX_Inductive_Table.docx"
doc.save(output_path)

output_path
